import streamlit as st
import os
import re
import json
import io
import base64
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from openai import OpenAI

try:
    from lunar_python import Solar, Lunar
    HAS_LUNAR = True
except ImportError:
    HAS_LUNAR = False

CONFIG_FILE = "naming_config.json"

# ================= 全局物理配置读写 (解决前后端数据隔离) =================
def load_global_config():
    default_config = {
        "api_base": "https://api.openai.com/v1",
        "api_key": "",
        "model_name": "gpt-4o",
        "usage_count": 0
    }
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                default_config.update(json.load(f))
        except:
            pass
    return default_config

def save_global_config(config_dict):
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(config_dict, f, ensure_ascii=False, indent=4)
    except:
        pass

# 启动时读取全局配置
global_config = load_global_config()

# ================= 页面基础设置 =================
st.set_page_config(page_title="名正言顺", layout="centered", page_icon="🌟")

# 注入自定义 CSS：强制覆盖生效
st.markdown("""
<style>
/* 1. 大标题缩小并改为金色 */
.main-title {
    font-size: 26px !important;
    color: #DAA520 !important; /* 金黄色 */
    font-weight: bold;
    text-align: center;
    margin-bottom: 0px;
}
.sub-title-desc {
    text-align: center;
    color: gray;
    font-size: 14px;
    margin-bottom: 20px;
}
/* 2. 小标题缩小并改为蓝色 */
.sub-title {
    font-size: 18px !important;
    color: #1E90FF !important; /* 亮蓝色 */
    font-weight: bold;
    margin-top: 15px;
    margin-bottom: 15px;
    border-left: 4px solid #1E90FF;
    padding-left: 8px;
}
/* 3. 多选框(风格偏好)强制改为优雅的绿底色 */
div[data-baseweb="select"] span[data-baseweb="tag"] {
    background-color: #2E8B57 !important; /* 海洋绿 */
    border: none !important;
}
div[data-baseweb="select"] span[data-baseweb="tag"] span {
    color: white !important;
    font-weight: bold;
}
div[data-baseweb="select"] span[data-baseweb="tag"] svg {
    fill: white !important;
}
</style>
""", unsafe_allow_html=True)

# ================= 会话状态初始化 =================
if 'admin_logged_in' not in st.session_state:
    st.session_state.admin_logged_in = False

# ================= 独立后台通道 (隐藏暗门) =================
is_admin_mode = st.query_params.get("admin") == "true"

if is_admin_mode:
    st.title("⚙️ 系统后台管理中心")
    if not st.session_state.admin_logged_in:
        pwd = st.text_input("请输入管理员密码", type="password")
        if st.button("登录后台", type="primary"):
            if pwd == "888888":  
                st.session_state.admin_logged_in = True
                st.success("登录成功！")
                st.rerun()
            else:
                st.error("密码错误！")
    else:
        st.success("✅ 管理员已授权登录")
        
        st.markdown("### 📊 本次运行实例数据")
        st.metric(label="累计成功推演生成次数", value=global_config.get("usage_count", 0))
        st.divider()
        
        st.markdown("### 🤖 AI 接口统一配置")
        
        # 1. 接口地址(Base URL) 下拉菜单（集成云端与主流本地 AI 服务类型）
        common_bases = [
            # 常见云端商用接口
            "https://api.openai.com/v1",
            "https://api.deepseek.com/v1",
            "https://dashscope.aliyuncs.com/compatible-mode/v1",
            "https://api.moonshot.cn/v1",
            # 本地常用 AI 运行框架及对应端口
            "http://localhost:11434/v1",  # Ollama 本地默认
            "http://localhost:1234/v1",   # LM Studio 本地默认
            "http://localhost:8000/v1",   # vLLM / FastChat 本地默认
            "http://localhost:9997/v1",   # Xinference 本地默认
            "http://localhost:5000/v1",   # Text-Gen-WebUI 本地默认
            "http://localhost:8080/v1",   # LocalAI 本地默认
            "自定义其他地址..."
        ]
        curr_base = global_config.get("api_base", "https://api.openai.com/v1")
        try:
            b_idx = common_bases.index(curr_base)
        except ValueError:
            b_idx = len(common_bases) - 1
            
        selected_base = st.selectbox("选择接口地址 (Base URL)", common_bases, index=b_idx)
        if selected_base == "自定义其他地址...":
            new_api_base = st.text_input("手动输入接口地址", value=curr_base)
        else:
            new_api_base = selected_base

        new_api_key = st.text_input("API 密钥 (API Key, 本地AI可为空)", value=global_config.get("api_key", ""), type="password")
        
        # 2. 模型名称 下拉菜单（集成云端与主流本地开源大模型规格）
        common_models = [
            # 云端常用模型
            "gpt-4o", "gpt-4o-mini", "gpt-3.5-turbo", 
            "deepseek-chat", "deepseek-reasoner", 
            "qwen-plus", "qwen-max", 
            "moonshot-v1-8k", 
            # 本地主流大模型 (兼容 Ollama / LM Studio 等命名)
            "deepseek-r1:7b", "deepseek-r1:8b", "deepseek-r1:14b", "deepseek-r1:32b",
            "deepseek-coder-v2", 
            "qwen3:30b", "qwen2.5:7b", "qwen2.5:14b", "qwen2.5:32b", "qwen2.5-7b-instruct",
            "llama3.1:8b", "llama3.2:3b", "llama3.3:70b",
            "glm4", "glm-4:9b", 
            "gemma3:27b", "gemma2:9b", 
            "mistral:7b",
            "自定义其他模型..."
        ]
        curr_model = global_config.get("model_name", "gpt-4o")
        try:
            m_idx = common_models.index(curr_model)
        except ValueError:
            m_idx = len(common_models) - 1
            
        selected_model = st.selectbox("选择模型名称 (Model)", common_models, index=m_idx)
        if selected_model == "自定义其他模型...":
            new_model_name = st.text_input("手动输入模型名称", value=curr_model)
        else:
            new_model_name = selected_model
        
        st.divider()
        if st.button("保存配置并退出后台", type="primary"):
            global_config["api_base"] = new_api_base
            global_config["api_key"] = new_api_key
            global_config["model_name"] = new_model_name
            save_global_config(global_config)
            
            st.session_state.admin_logged_in = False
            st.success("配置已成功保存！请去掉网址后缀的 ?admin=true 返回前台。")
            
    st.stop()


# ================= 前端主界面 (用户端纯净版) =================
st.markdown('<div class="main-title">🌟 名正言顺 - 国学起名系统 v3.0</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title-desc">融合传统周易八卦、五行生克与现代 AI 智能的前沿起名引擎。</div>', unsafe_allow_html=True)

# 1. 基础信息填写
st.markdown('<div class="sub-title">1. 基础信息填写</div>', unsafe_allow_html=True)
mode = st.radio("选择起名类型", ["为人起名", "为公司起名"], horizontal=True)

col1, col2 = st.columns(2)
with col1:
    cal_type = st.selectbox("历法选择", ["公历", "农历"])
    year = st.number_input("出生年", min_value=1920, max_value=2040, value=1990)
    month = st.number_input("出生月", min_value=1, max_value=12, value=1)
with col2:
    day = st.number_input("出生日", min_value=1, max_value=31, value=1)
    hour = st.selectbox("时辰", [f"{h} 时" for h in range(24)] + ["不清楚时辰"], index=12)
    gender = st.radio("性别/属性", ["男 (乾造)", "女 (坤造)"], horizontal=True)

# 专属字段
surname, birth_place = "", ""
comp_prefix, comp_suffix, comp_industry = "", "", ""

if mode == "为人起名":
    c1, c2 = st.columns(2)
    with c1: surname = st.text_input("姓氏 (必填)", value="李")
    with c2: birth_place = st.text_input("出生地 (选填)")
else:
    c1, c2 = st.columns(2)
    with c1: comp_prefix = st.text_input("公司前缀 (如: XX市)")
    with c2: comp_suffix = st.text_input("公司后缀 (如: 科技有限公司)")
    comp_industry = st.text_input("行业/业务范围", value="互联网科技")

# 2. 风格偏好与要求
st.markdown('<div class="sub-title">2. 风格偏好与输出要求</div>', unsafe_allow_html=True)
all_prefs = ["寓意好", "音韵美", "合五行", "合数理", "有典故", "有个性", "少歧义", "求时尚", "中性别", "地域性", "中英文"]
selected_prefs = st.multiselect("风格偏好 (可多选)", all_prefs, default=["寓意好", "合五行", "合数理", "音韵美", "有典故"])

c3, c4 = st.columns(2)
with c3: name_length = st.text_input("名字字数要求", value="3字")
with c4: name_count = st.number_input("生成方案个数", min_value=1, value=5)
other_req = st.text_area("其他补充要求 (如希望名字大气、避免生僻字;  特定的字辈、避免的字等)")


# ================= 核心推演逻辑 =================
def get_bazi_info():
    gender_str = "乾造" if "男" in gender else "坤造"
    bazi_label = "法人八字" if mode == "为公司起名" else "命主八字"
    
    if not HAS_LUNAR:
        return f"公历时间：{cal_type} {year}年{month}月{day}日\n{bazi_label}：缺 (未安装排盘库) ({gender_str})"
    
    try:
        h = int(hour.split(" ")[0]) if hour != "不清楚时辰" else 12
        if cal_type == "公历":
            solar = Solar.fromYmdHms(int(year), int(month), int(day), h, 0, 0)
            lunar = solar.getLunar()
        else:
            lunar = Lunar.fromYmdHms(int(year), int(month), int(day), h, 0, 0)
            solar = lunar.getSolar()
            
        ba_zi = lunar.getEightChar()
        return f"出生公历：{solar.getYear()}年{solar.getMonth()}月{solar.getDay()}日 {h}时\n出生农历：{lunar.getYear()}年{lunar.getMonthInChinese()}月{lunar.getDayInChinese()} {h}时\n{bazi_label}：{ba_zi.getYear()} {ba_zi.getMonth()} {ba_zi.getDay()} {ba_zi.getTime()}（{gender_str}）"
    except Exception:
        return f"日期解析异常 ({gender_str})"


def generate_word_doc(content, bazi_info):
    doc = Document()
    # 统一全局字体为更高级的微软雅黑，11磅
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)

    fav_element, taboo_element = "火", "水"
    color_match = re.search(r'【五行色彩：喜=(.)，忌=(.)】', content)
    if color_match:
        fav_element, taboo_element = color_match.group(1), color_match.group(2)
        content = content.replace(color_match.group(0), '').strip()

    color_map = {'木': RGBColor(0, 128, 0), '火': RGBColor(255, 0, 0), '土': RGBColor(255, 255, 0), '金': RGBColor(255, 255, 255), '水': RGBColor(0, 0, 0)}
    hex_map = {'木': 'E8F5E9', '火': 'FFEBEE', '土': 'FFF9C4', '金': 'F5F5F5', '水': 'E0E0E0'}

    # 封面大标题
    title_text = "个人专属起名策划方案" if mode == "为人起名" else "公司品牌起名策划方案"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(28)
    title_para.paragraph_format.space_after = Pt(5)
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    pPr = title_para._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_map.get(taboo_element, 'FFFFFF'))
    pPr.append(shd)

    title_run = title_para.add_run(title_text)
    title_run.font.name = 'SimHei'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = color_map.get(fav_element, RGBColor(0, 0, 0))

    time_para = doc.add_paragraph(f"测算时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    time_para.runs[0].font.size = Pt(10)
    time_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # 核心排版引擎升级：美观高档
    for line in content.split('\n'):
        line = line.strip()
        line = re.sub(r'<[^>]+>', '', line) # 清洗 HTML 标签
        if not line: continue
        
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5 # 舒适的1.5倍行距
        
        if line.startswith("第一部分") or line.startswith("第二部分"):
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(44, 62, 80) # 高级深沉蓝灰色
            
            if "第一部分" in line:
                # 给八字加上高级的浅褐色底纹边框效果
                doc.add_paragraph("")
                bazi_para = doc.add_paragraph(bazi_info)
                bazi_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                bazi_para.paragraph_format.space_before = Pt(10)
                bazi_para.paragraph_format.space_after = Pt(10)
                
                bazi_pPr = bazi_para._element.get_or_add_pPr()
                bazi_shd = OxmlElement('w:shd')
                bazi_shd.set(qn('w:val'), 'clear')
                bazi_shd.set(qn('w:color'), 'auto')
                bazi_shd.set(qn('w:fill'), 'F9EBEA') # 典雅浅褐色背景
                bazi_pPr.append(bazi_shd)

                run_bazi = bazi_para.runs[0]
                run_bazi.font.size = Pt(12)
                run_bazi.font.bold = True
                run_bazi.font.color.rgb = RGBColor(139, 69, 19) # 马鞍棕色
                doc.add_paragraph("")
                
        elif line.startswith("【") and "方案" in line:
            # 方案大标题
            p.paragraph_format.space_before = Pt(15)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(line)
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(178, 34, 34) # 砖红色，极具国学高级感
            
        elif re.match(r'^\d+\.', line):
            # 类似 "1. 拼音：" 的主列表：缩进、加粗标头
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'：|:', line, 1)
            if len(parts) == 2:
                r1 = p.add_run(parts[0] + "：")
                r1.font.bold = True
                r1.font.color.rgb = RGBColor(39, 174, 96) # 典雅翡翠绿
                r2 = p.add_run(parts[1])
                r2.font.color.rgb = RGBColor(50, 50, 50) # 深灰色代替纯黑，更柔和
            else:
                r = p.add_run(line)
                r.font.color.rgb = RGBColor(50, 50, 50)
                
        elif line.startswith("- "):
            # 类似 "- 核心名数理分析：" 的子列表：二级阶梯缩进
            p.paragraph_format.left_indent = Pt(28)
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'：|:', line, 1)
            if len(parts) == 2:
                r1 = p.add_run(parts[0] + "：")
                r1.font.bold = True
                r1.font.color.rgb = RGBColor(139, 69, 19) # 次级强调色：古铜色
                r2 = p.add_run(parts[1])
                r2.font.color.rgb = RGBColor(60, 60, 60)
            else:
                r = p.add_run(line)
                r.font.color.rgb = RGBColor(60, 60, 60)
                
        else:
            # 普通正文
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line)
            r.font.color.rgb = RGBColor(60, 60, 60)

    disclaimer = doc.add_paragraph("\n[ 注：本起名策划方案系依据传统国学经典推演，仅供参考 ]")
    disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    disclaimer.runs[0].font.size = Pt(10)
    disclaimer.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio


# 生成按钮
st.divider()
if st.button("🚀 开始启动", type="primary", use_container_width=True):
    is_local_ai = "localhost" in global_config.get("api_base", "") or "127.0.0.1" in global_config.get("api_base", "")
    
    if not global_config.get("api_key") and not is_local_ai:
        st.warning("⚠️ 接口暂未打通，请联系管理员配置。")
    else:
        with st.spinner("算力运转中，正在为您结合八字与国学全力推演，请稍候..."):
            try:
                bazi_info = get_bazi_info()
                
                # 构造 Prompt
                prompt = "你是一个精通中国传统文化、周易八卦、五行生克、三才五格数理、诗词歌赋的资深国学起名大师。\n\n"
                
                prompt += f"【背景信息】\n- 性别：{gender}\n- 出生时间与八字：\n{bazi_info}\n"
                prompt += "（请在第一部分采用最正宗、最权威、最精确的传统子平八字命理体系，综合日元旺衰、格局分析、调候通关等，给出深度专业的五行用神、忌神喜忌分析）\n"

                num_match = re.search(r'\d+', name_length)
                n_total = int(num_match.group()) if num_match else 3
                
                if mode == "为人起名":
                    prompt += f"- 起名类型：个人起名\n- 姓氏：{surname}\n"
                    if birth_place: prompt += f"- 出生地：{birth_place}\n"
                    
                    surname_len = len(surname) if surname else 1
                    n_names = max(1, n_total - surname_len)
                    mask_chars = "〇" * n_names
                    template_mask = f"{surname}{mask_chars}"
                    
                    prompt += f"\n【⚠️ 严密字数核心铁律 - 绝不允许出错！】\n"
                    prompt += f"1. 包含姓氏在内，全名总字数必须严格等于 {n_total} 个汉字！\n"
                    prompt += f"2. 因为姓氏“{surname}”占 {surname_len} 个字，所以你起的名字（不含姓氏）只能是 {n_names} 个字！\n"
                    prompt += f"3. 强制填空：名字必须完美填入模板【 {template_mask} 】（将〇替换为你起的字），多一个字、少一个字都将被判定为失败。\n"
                else:
                    prompt += f"- 起名类型：公司/品牌起名\n- 行业及经营范围：{comp_industry}\n"
                    if comp_prefix or comp_suffix: 
                        prompt += f"- 结构：{comp_prefix} + [核心商号] + {comp_suffix}\n"
                        prompt += "【极其重要】：在公司起名中，必须同时独立测算“核心名”和“公司全称”的数理吉凶！不能混为一谈。\n"
                    prompt += f"\n【⚠️ 核心铁律】\n1. 核心商号必须绝对等于 {n_total} 个汉字！\n2. 强制填空：必须完全填入模板【 {'〇' * n_total} 】。\n"

                prompt += f"【偏好要求】\n- 勾选要求：{', '.join(selected_prefs)}\n- 补充要求：{other_req}\n"
                prompt += "\n【五行颜色底层指令提取】\n请在整篇回答的最开头（第一行），原封不动输出：\n【五行色彩：喜=X，忌=Y】（X和Y替换为单字，勿带其他文字）。\n"
                
                prompt += "\n【排版与视觉美学指令】\n"
                prompt += "为了使输出结果在手机和电脑端网页显示和谐美观，请严格使用 Markdown 和 HTML 标签进行排版优化：\n"
                prompt += "1. 请使用 <span style='color: #特定颜色; font-weight: bold;'> 标签包裹关键属性（如五行、分数等）。\n"
                prompt += "2. 颜色匹配原则：木用绿色(#2E8B57)，火用红色(#B22222)，土用棕黄(#B8860B)，金用金色(#DAA520)，水用蓝色(#1E90FF)。\n"
                prompt += "3. 长段落排版要层次分明、采用无序列表或区块引用。\n"
                
                # 【强化版】：强制防止断流，确保方案个数 100% 达标
                prompt += f"\n【⚠️ 特别警告：完整性要求】\n"
                prompt += f"本次请求生成的方案数量为 {name_count} 个。请你务必分配好输出篇幅，**绝对保证从“方案一”完整输出到“方案{name_count}”**，绝不允许中途截断或省略！如果内容较多，请精简“典故”和“寓意”的字数，但必须保证总数达到 {name_count} 个。\n\n"
                
                prompt += "必须严格按以下子标题结构输出（不可省略序号）：\n"
                prompt += "第一部分：五行喜忌分析\n"
                prompt += "第二部分：起名方案\n"
                prompt += "【方案一：XXX】\n"
                prompt += "1. 拼音：\n"
                prompt += "2. 五行：\n"
                
                if mode == "为公司起名" and (comp_prefix or comp_suffix):
                    prompt += "3. 数理分析（必须分别独立作答）：\n"
                    prompt += "   - 核心名数理分析（仅测算核心名的笔画吉凶与评分）：\n"
                    prompt += f"   - 公司全称数理分析（整体测算“{comp_prefix} + 核心名 + {comp_suffix}”的笔画吉凶与评分）：\n"
                else:
                    prompt += "3. 数理：\n"
                    
                prompt += "4. 音韵：\n"
                prompt += "5. 典故：\n"
                prompt += "6. 寓意：\n"

                actual_api_key = global_config["api_key"] if global_config["api_key"] else "sk-local-dummy-key"
                
                client = OpenAI(api_key=actual_api_key, base_url=global_config["api_base"])
                response = client.chat.completions.create(
                    model=global_config["model_name"],
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=8192
                )
                
                ai_content = response.choices[0].message.content
                
                st.success("✅ 推演成功！")
                
                global_config["usage_count"] = global_config.get("usage_count", 0) + 1
                save_global_config(global_config)
                
                clean_html_content = ai_content.replace(re.search(r'【五行色彩：喜=.*，忌=.*】', ai_content).group(0), '') if re.search(r'【五行色彩：喜=.*，忌=.*】', ai_content) else ai_content
                st.markdown(clean_html_content, unsafe_allow_html=True)

                st.info("💡 提示：如遇部分手机浏览器点击下方下载按钮无反应，请使用下方的【📲 备用直连下载通道】。")
                
                word_file = generate_word_doc(ai_content, bazi_info)
                filename = f"起名策划方案_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                
                st.download_button(
                    label="📥 保存为 Word 策划方案文档 (推荐)",
                    data=word_file.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                b64 = base64.b64encode(word_file.getvalue()).decode()
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" style="display: block; width: 100%; text-align: center; padding: 12px; margin-top: 5px; background-color: #f0f2f6; color: #1E90FF; border-radius: 8px; text-decoration: none; font-weight: bold; border: 1px solid #d5dce5; box-shadow: 0 2px 4px rgba(0,0,0,0.05);">📲 备用直连下载通道（专供手机端）</a>'
                st.markdown(href, unsafe_allow_html=True)
                
            except Exception as e:
                st.error(f"❌ 生成失败，请联系管理员检查配置或网络。详细报错：{e}")
